"""
AltoLex — app_v4.py
Full implementation:
  - Client + case management (saved to DB)
  - Document upload embedded via voyage-law-2, stored in documents table
  - Documents tied to a specific case OR firm-wide (ordinances)
  - AI Search with RAG over stored documents
  - Admin panel for user management
"""

import streamlit as st
import anthropic
import os, io, base64, hashlib, json, re
from pathlib import Path
from datetime import date
from auth import (login, logout, validate_session,
                  create_user, update_user, deactivate_user,
                  reactivate_user, list_users, log_action)
from db import ScopedDB, raw_client

try:
    import fitz;        PYMUPDF_OK = True
except ImportError:     PYMUPDF_OK = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:     DOCX_OK = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    REPORTLAB_OK = True
except ImportError:     REPORTLAB_OK = False
try:
    import PyPDF2;      PYPDF2_OK = True
except ImportError:     PYPDF2_OK = False
try:
    import voyageai;    VOYAGE_OK = True
except ImportError:     VOYAGE_OK = False


st.set_page_config(page_title="AltoLex", page_icon="⚖️", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600&family=DM+Sans:wght@300;400;500&display=swap');

/* ── Font ── */
html, body, [class*="css"], .stMarkdown, .stText,
p, li, span, div { font-family: 'DM Sans', sans-serif; }
h1, h2, h3 { font-family: 'Playfair Display', serif !important; }

/* ── Use Streamlit's own theme tokens so both light and dark work ── */
.stApp { background: var(--background-color); }

/* All text inherits Streamlit theme colour — never hardcode dark */
.stMarkdown, .stMarkdown p, .stMarkdown li,
.stMarkdown span, .stText { color: var(--text-color) !important; }

/* Expander headers */
.streamlit-expanderHeader { color: var(--text-color) !important; }

/* Badges — coloured borders/text only, transparent background */
.badge {
    display:inline-block; padding:3px 10px; border-radius:20px;
    font-size:11px; font-weight:600; letter-spacing:0.06em;
    text-transform:uppercase; margin-right:4px;
}
.badge-firm   { background:rgba(184,147,90,0.15); color:#c9a96e; border:1px solid rgba(184,147,90,0.4); }
.badge-case   { background:rgba(80,130,220,0.15); color:#7aabff; border:1px solid rgba(80,130,220,0.4); }
.badge-common { background:rgba(60,160,100,0.15); color:#6dcf96; border:1px solid rgba(60,160,100,0.4); }
.badge-admin  { background:rgba(200,60,80,0.12);  color:#e07080; border:1px solid rgba(200,60,80,0.3); }

/* Disclaimer box — adapts to theme */
.disclaimer {
    background: rgba(184,147,90,0.08);
    border: 1px solid rgba(184,147,90,0.25);
    border-radius:8px; padding:12px 16px;
    font-size:12px; line-height:1.6;
    color: var(--text-color);
    opacity: 0.8;
}

/* Inline small/caption text — use opacity instead of fixed colour */
small { opacity: 0.65; }

/* Card-style containers */
.card {
    background: var(--secondary-background-color);
    border: 1px solid rgba(184,147,90,0.2);
    border-radius:10px; padding:16px 18px; margin-bottom:10px;
}

/* Login card */
.login-card {
    background: var(--secondary-background-color);
    border: 1px solid rgba(184,147,90,0.25);
    border-radius:12px; padding:32px;
    max-width:420px; margin:60px auto 24px;
}
.login-card .wordmark {
    font-family:'Playfair Display',serif;
    font-size:28px; font-weight:600;
    color: var(--text-color);
    margin-bottom:3px;
}
.login-card .tagline {
    font-size:10px; text-transform:uppercase;
    letter-spacing:0.1em; opacity:0.5;
    margin-bottom:28px;
}
</style>
""", unsafe_allow_html=True)

SYSTEM_PROMPT = """You are AltoLex, a professional legal information assistant for a law firm in Sri Lanka.
Retrieved context from the firm document library is below - prefer this over general knowledge.
Cite source document names when referencing retrieved content.
RULES:
1. Never give definitive legal advice - provide legal information only
2. Always recommend attorney review for specific situations
3. If context lacks relevant information, say so clearly - never fabricate
4. Flag urgent deadlines or limitation periods prominently
5. End substantive responses with: "This is legal information only. Please consult a qualified attorney."
"""

ROLES        = ["admin", "partner", "associate", "paralegal", "readonly"]
AREAS_OF_LAW = ["Contract dispute", "Property / conveyancing", "Employment", "Family law",
                "Commercial / corporate", "Intellectual property", "Criminal defence", "Other"]
CHUNK_SIZE   = 500
CHUNK_OVERLAP = 80


# ---- API clients -------------------------------------------------------
@st.cache_resource
def get_anthropic():
    key = os.environ.get("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY", "")
    return anthropic.Anthropic(api_key=key)

@st.cache_resource
def get_voyage():
    key = os.environ.get("VOYAGE_API_KEY") or st.secrets.get("VOYAGE_API_KEY", "")
    return voyageai.Client(api_key=key)


# ---- Session helpers ---------------------------------------------------
def ctx() -> dict:
    return st.session_state.get("ctx", {})

def sdb() -> ScopedDB:
    return ScopedDB(ctx()["firm_id"])

def is_admin()    -> bool: return ctx().get("role") == "admin"
def is_readonly() -> bool: return ctx().get("role") == "readonly"


# ---- Client / case DB helpers -----------------------------------------
def get_clients() -> list:
    r = sdb().table("clients").select("id,full_name,email,phone,created_at").order("full_name").execute()
    return r.data or []

def get_cases(client_id: str) -> list:
    r = sdb().table("cases").select("id,title,area_of_law,status,created_at") \
             .eq("client_id", client_id).order("created_at", desc=True).execute()
    return r.data or []

def create_client_record(full_name: str, email: str, phone: str) -> dict:
    r = sdb().table("clients").insert({"full_name": full_name, "email": email or "", "phone": phone or ""}).execute()
    return r.data[0]

def create_case_record(client_id: str, title: str, area: str) -> dict:
    r = sdb().table("cases").insert({"client_id": client_id, "title": title, "area_of_law": area, "status": "open"}).execute()
    return r.data[0]

def get_stored_docs(case_id: str = None, firm_wide: bool = False) -> list:
    if firm_wide:
        rc = raw_client()
        rows = rc.table("documents").select("metadata,file_hash,created_at") \
                 .eq("firm_id", ctx()["firm_id"]).is_("case_id", "null").execute().data or []
    else:
        rows = sdb().table("documents").select("metadata,file_hash,created_at") \
                    .eq("case_id", case_id).execute().data or []
    seen = set(); result = []
    for row in rows:
        h = row.get("file_hash", "")
        if h not in seen:
            seen.add(h)
            result.append({"name": row["metadata"].get("source", "unknown"),
                           "file_hash": h, "created_at": row.get("created_at", "")})
    return result


# ---- Document extraction ----------------------------------------------
def extract_text(file_bytes: bytes, filename: str) -> tuple:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" and PYPDF2_OK:
        try:
            pages = [p.extract_text() or "" for p in PyPDF2.PdfReader(io.BytesIO(file_bytes)).pages]
            t = "\n\n".join(pages).strip()
            if len(t) > 100:
                return t, "text"
        except Exception:
            pass
        return "", "vision_needed"
    elif ext in [".docx", ".doc"] and DOCX_OK:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append("## " + t if p.style.name.startswith("Heading") else t)
            for table in doc.tables:
                parts.append("[TABLE]")
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
                parts.append("[END TABLE]")
            return "\n".join(parts), "docx"
        except Exception:
            pass
    return "", "empty"

def chunk_text(text: str) -> list:
    words = text.split()
    step  = CHUNK_SIZE - CHUNK_OVERLAP
    return [" ".join(words[i:i+CHUNK_SIZE]).strip()
            for i in range(0, len(words), step) if words[i:i+CHUNK_SIZE]]

def doc_type_tag(name: str) -> str:
    n = name.lower()
    if any(k in n for k in ["nda","non-disclosure"]):          return "nda"
    if any(k in n for k in ["employment","service"]):          return "employment"
    if any(k in n for k in ["lease","tenancy","rental"]):      return "property"
    if any(k in n for k in ["judgment","ruling","case"]):      return "case_law"
    if any(k in n for k in ["ordinance","act","regulation",
                             "statute","gazette"]):             return "ordinance"
    return "general"

def embed_and_store(file_bytes: bytes, filename: str, firm_id: str, case_id: str = None) -> dict:
    """Extract -> chunk -> embed -> store. case_id=None = firm-wide."""
    fhash = hashlib.md5(file_bytes).hexdigest()
    existing = raw_client().table("documents").select("id") \
                .eq("firm_id", firm_id).eq("file_hash", fhash).limit(1).execute()
    if existing.data:
        return {"skipped": True, "chunks": 0, "method": ""}

    text, method = extract_text(file_bytes, filename)

    # OCR for scanned PDFs via Claude vision
    if method == "vision_needed" and PYMUPDF_OK:
        try:
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            mat = fitz.Matrix(150/72, 150/72)
            page_texts = []
            for i, page in enumerate(pdf_doc):
                if i >= 15: break
                pix = page.get_pixmap(matrix=mat)
                img_b64 = base64.standard_b64encode(pix.tobytes("png")).decode()
                resp = get_anthropic().messages.create(
                    model="claude-sonnet-4-20250514", max_tokens=2000,
                    messages=[{"role": "user", "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img_b64}},
                        {"type": "text", "text": "Extract all text from this document page. Return only the text."}
                    ]}])
                page_texts.append(resp.content[0].text)
            text = "\n\n".join(page_texts)
            method = "vision_ocr"
        except Exception as e:
            return {"skipped": False, "chunks": 0, "method": "error", "error": str(e)}

    if not text.strip():
        return {"skipped": False, "chunks": 0, "method": "empty"}

    chunks  = chunk_text(text)
    vo      = get_voyage()
    vectors = []
    for i in range(0, len(chunks), 64):
        r = vo.embed(chunks[i:i+64], model="voyage-law-2", input_type="document")
        vectors.extend(r.embeddings)

    # Upload to Supabase Storage for later preview retrieval
    storage_path = storage_upload(file_bytes, firm_id, case_id, filename)

    rows = [{"firm_id": firm_id, "case_id": case_id, "content": chunk, "embedding": vector,
             "metadata": {"source": filename, "doc_type": doc_type_tag(filename),
                          "chunk_idx": idx, "storage_path": storage_path},
             "file_hash": fhash}
            for idx, (chunk, vector) in enumerate(zip(chunks, vectors))]

    rc = raw_client()
    for i in range(0, len(rows), 50):
        rc.table("documents").insert(rows[i:i+50]).execute()

    return {"skipped": False, "chunks": len(chunks), "method": method, "storage_path": storage_path}




# ---- Supabase Storage helpers ------------------------------------------
STORAGE_BUCKET = "altolex-documents"

def storage_upload(file_bytes: bytes, firm_id: str, case_id: str, filename: str) -> str:
    """Upload file bytes to Supabase Storage. Returns storage path."""
    import mimetypes
    rc = raw_client()
    safe_name = re.sub(r"[^\w.\-]", "_", filename)
    folder = f"{firm_id}/{case_id or 'common'}"
    path   = f"{folder}/{safe_name}"
    mime   = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    try:
        rc.storage.from_(STORAGE_BUCKET).upload(
            path=path,
            file=file_bytes,
            file_options={"content-type": mime, "upsert": "true"}
        )
    except Exception as e:
        # If bucket doesn't exist yet, create it then retry
        if "not found" in str(e).lower() or "does not exist" in str(e).lower():
            try:
                rc.storage.create_bucket(STORAGE_BUCKET, options={"public": False})
                rc.storage.from_(STORAGE_BUCKET).upload(
                    path=path, file=file_bytes,
                    file_options={"content-type": mime, "upsert": "true"}
                )
            except Exception:
                return ""
        else:
            return ""
    return path

def storage_download(storage_path: str) -> bytes | None:
    """Download file bytes from Supabase Storage."""
    try:
        rc  = raw_client()
        res = rc.storage.from_(STORAGE_BUCKET).download(storage_path)
        return res
    except Exception:
        return None

def get_stored_docs_with_path(case_id: str = None, firm_wide: bool = False) -> list:
    """List unique docs with their storage paths for preview."""
    if firm_wide:
        rows = raw_client().table("documents") \
                .select("metadata,file_hash,created_at") \
                .eq("firm_id", ctx()["firm_id"]).is_("case_id", "null").execute().data or []
    else:
        rows = sdb().table("documents").select("metadata,file_hash,created_at") \
                    .eq("case_id", case_id).execute().data or []
    seen = set(); result = []
    for row in rows:
        h = row.get("file_hash", "")
        if h not in seen:
            seen.add(h)
            meta = row.get("metadata", {})
            result.append({
                "name":         meta.get("source", "unknown"),
                "file_hash":    h,
                "storage_path": meta.get("storage_path", ""),
                "created_at":   row.get("created_at", ""),
                "doc_type":     meta.get("doc_type", "general"),
            })
    return result


# ---- PDF viewer HTML component -----------------------------------------
def render_pdf_viewer(file_bytes: bytes, filename: str, height: int = 600):
    """Render a PDF inline using PDF.js via Streamlit HTML component."""
    if not file_bytes:
        st.warning("Could not load document.")
        return
    b64 = base64.b64encode(file_bytes).decode()
    html = f"""
    <div id="pdf-outer" style="border:1px solid var(--border-color,#ddd);border-radius:8px;overflow:hidden;background:#2a2a2a;">

      <!-- Toolbar -->
      <div style="display:flex;align-items:center;gap:8px;padding:8px 14px;
                  background:#1a1a2e;color:#d4aa7a;font-size:13px;font-family:sans-serif;">
        <span style="flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap">{filename}</span>
        <button onclick="changePage(-1)" style="background:rgba(255,255,255,0.1);border:none;color:#d4aa7a;
          padding:4px 10px;border-radius:5px;cursor:pointer;font-size:13px">&#8249;</button>
        <span id="page-info" style="font-size:12px;white-space:nowrap">Page 1</span>
        <button onclick="changePage(1)"  style="background:rgba(255,255,255,0.1);border:none;color:#d4aa7a;
          padding:4px 10px;border-radius:5px;cursor:pointer;font-size:13px">&#8250;</button>
        <button onclick="zoomIn()"  style="background:rgba(255,255,255,0.1);border:none;color:#d4aa7a;
          padding:4px 8px;border-radius:5px;cursor:pointer">+</button>
        <button onclick="zoomOut()" style="background:rgba(255,255,255,0.1);border:none;color:#d4aa7a;
          padding:4px 8px;border-radius:5px;cursor:pointer">-</button>
        <button onclick="printPDF()" style="background:rgba(184,147,90,0.2);border:1px solid rgba(184,147,90,0.4);
          color:#d4aa7a;padding:4px 12px;border-radius:5px;cursor:pointer;font-size:12px">Print</button>
      </div>

      <!-- Canvas -->
      <div style="overflow:auto;max-height:{height}px;display:flex;justify-content:center;padding:16px;">
        <canvas id="pdf-canvas" style="box-shadow:0 4px 20px rgba(0,0,0,0.5);border-radius:3px;"></canvas>
      </div>
    </div>

    <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
    <script>
      pdfjsLib.GlobalWorkerOptions.workerSrc =
        'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';

      const pdfData  = atob('{b64}');
      const pdfBytes = new Uint8Array(pdfData.length);
      for (let i=0; i<pdfData.length; i++) pdfBytes[i] = pdfData.charCodeAt(i);

      let pdfDoc=null, curPage=1, scale=1.4;

      pdfjsLib.getDocument({{data:pdfBytes}}).promise.then(doc => {{
        pdfDoc = doc;
        renderPage(1);
      }});

      function renderPage(n) {{
        pdfDoc.getPage(n).then(page => {{
          const vp  = page.getViewport({{scale}});
          const canvas = document.getElementById('pdf-canvas');
          const ctx  = canvas.getContext('2d');
          canvas.width  = vp.width;
          canvas.height = vp.height;
          page.render({{canvasContext:ctx, viewport:vp}});
          document.getElementById('page-info').textContent =
            'Page ' + curPage + ' of ' + pdfDoc.numPages;
        }});
      }}

      function changePage(d) {{
        const n = curPage + d;
        if (!pdfDoc || n < 1 || n > pdfDoc.numPages) return;
        curPage = n; renderPage(curPage);
      }}

      function zoomIn()  {{ scale = Math.min(3.0, scale+0.25); if(pdfDoc) renderPage(curPage); }}
      function zoomOut() {{ scale = Math.max(0.5, scale-0.25); if(pdfDoc) renderPage(curPage); }}

      function printPDF() {{
        const win = window.open('');
        win.document.write('<html><body style="margin:0">');
        win.document.write('<canvas id="pc"></canvas>');
        win.document.write('</body></html>');
        win.document.close();
        const allPages = async () => {{
          for (let i=1; i<=pdfDoc.numPages; i++) {{
            const pg = await pdfDoc.getPage(i);
            const vp = pg.getViewport({{scale:1.5}});
            const c  = win.document.getElementById('pc');
            c.width=vp.width; c.height=vp.height;
            await pg.render({{canvasContext:c.getContext('2d'),viewport:vp}}).promise;
          }}
          win.print();
        }};
        allPages();
      }}
    </script>
    """
    st.components.v1.html(html, height=height + 60, scrolling=False)


# ---- Contract templates -------------------------------------------------
CONTRACT_TEMPLATES = {
    "Non-Disclosure Agreement (NDA)": {
        "description": "Mutual or one-way confidentiality agreement between two parties.",
        "fields": [
            {"key": "disclosing_party",   "label": "Disclosing party name",        "type": "text"},
            {"key": "receiving_party",    "label": "Receiving party name",          "type": "text"},
            {"key": "purpose",            "label": "Purpose of disclosure",         "type": "text",
             "placeholder": "e.g. evaluating a potential business partnership"},
            {"key": "duration_years",     "label": "Confidentiality period (years)","type": "number", "default": 2},
            {"key": "governing_law",      "label": "Governing law / jurisdiction",  "type": "text",
             "default": "Sri Lanka"},
            {"key": "date",               "label": "Agreement date",                "type": "date"},
        ],
    },
    "Employment Contract": {
        "description": "Standard contract of employment for a permanent employee.",
        "fields": [
            {"key": "employer_name",      "label": "Employer name",                 "type": "text"},
            {"key": "employer_address",   "label": "Employer address",              "type": "text"},
            {"key": "employee_name",      "label": "Employee full name",            "type": "text"},
            {"key": "employee_nic",       "label": "Employee NIC number",           "type": "text"},
            {"key": "position",           "label": "Job title / position",          "type": "text"},
            {"key": "start_date",         "label": "Start date",                    "type": "date"},
            {"key": "basic_salary",       "label": "Basic salary (LKR/month)",      "type": "text"},
            {"key": "notice_period_months","label": "Notice period (months)",       "type": "number", "default": 3},
            {"key": "probation_months",   "label": "Probation period (months)",     "type": "number", "default": 6},
            {"key": "working_hours",      "label": "Working hours",                 "type": "text",
             "default": "8:30 AM to 5:30 PM, Monday to Friday"},
        ],
    },
    "Tenancy Agreement": {
        "description": "Residential or commercial lease agreement.",
        "fields": [
            {"key": "landlord_name",      "label": "Landlord full name",            "type": "text"},
            {"key": "tenant_name",        "label": "Tenant full name",              "type": "text"},
            {"key": "property_address",   "label": "Property address",              "type": "text"},
            {"key": "monthly_rent",       "label": "Monthly rent (LKR)",            "type": "text"},
            {"key": "deposit",            "label": "Security deposit (LKR)",        "type": "text"},
            {"key": "lease_start",        "label": "Lease start date",              "type": "date"},
            {"key": "lease_months",       "label": "Lease duration (months)",       "type": "number", "default": 12},
            {"key": "notice_days",        "label": "Notice period (days)",          "type": "number", "default": 30},
            {"key": "use",                "label": "Permitted use",                 "type": "text",
             "default": "residential purposes only"},
        ],
    },
    "Power of Attorney": {
        "description": "General or specific power of attorney under Sri Lankan law.",
        "fields": [
            {"key": "donor_name",         "label": "Donor (grantor) full name",     "type": "text"},
            {"key": "donor_nic",          "label": "Donor NIC",                     "type": "text"},
            {"key": "donor_address",      "label": "Donor address",                 "type": "text"},
            {"key": "attorney_name",      "label": "Attorney-in-fact full name",    "type": "text"},
            {"key": "attorney_nic",       "label": "Attorney-in-fact NIC",          "type": "text"},
            {"key": "powers",             "label": "Powers granted",                "type": "textarea",
             "placeholder": "e.g. to sell, transfer and convey the property at No. 42 Galle Road, Colombo"},
            {"key": "date",               "label": "Date of execution",             "type": "date"},
        ],
    },
    "Statutory Declaration": {
        "description": "Formal declaration of facts made under oath.",
        "fields": [
            {"key": "declarant_name",     "label": "Declarant full name",           "type": "text"},
            {"key": "declarant_nic",      "label": "Declarant NIC",                 "type": "text"},
            {"key": "declarant_address",  "label": "Declarant address",             "type": "text"},
            {"key": "declaration_facts",  "label": "Facts being declared",          "type": "textarea",
             "placeholder": "State the facts to be declared..."},
            {"key": "purpose",            "label": "Purpose of declaration",        "type": "text"},
            {"key": "date",               "label": "Date",                          "type": "date"},
        ],
    },
    "Service Agreement": {
        "description": "Agreement for professional or consulting services.",
        "fields": [
            {"key": "client_name",        "label": "Client name",                   "type": "text"},
            {"key": "provider_name",      "label": "Service provider name",         "type": "text"},
            {"key": "services",           "label": "Description of services",       "type": "textarea",
             "placeholder": "Describe the services to be provided..."},
            {"key": "fee",                "label": "Fee / payment terms",           "type": "text",
             "placeholder": "e.g. LKR 50,000 per month, payable on the 1st"},
            {"key": "start_date",         "label": "Start date",                    "type": "date"},
            {"key": "duration",           "label": "Duration",                      "type": "text",
             "default": "12 months"},
            {"key": "notice_days",        "label": "Termination notice (days)",     "type": "number", "default": 30},
            {"key": "governing_law",      "label": "Governing law",                 "type": "text",
             "default": "Sri Lanka"},
        ],
    },
}


def draft_contract_with_ai(template_name: str, variables: dict, firm_knowledge: str = "") -> str:
    """Call Claude to draft a full contract from template name and variables."""
    vars_text = "\n".join(f"  {k}: {v}" for k, v in variables.items() if v)
    prompt = f"""Draft a complete, professional {template_name} under Sri Lankan law.

Use the following details provided by the attorney:
{vars_text}

Requirements:
- Use formal legal language appropriate for Sri Lanka
- Include all standard clauses expected in a {template_name}
- Reference applicable Sri Lankan statutes where relevant
- Structure with clear numbered clauses and sub-clauses
- Include signature blocks for all parties at the end
- Mark any clause that needs attorney review with [REVIEW]
- Do NOT use placeholder text like [INSERT] — use the provided details throughout
- Today's date is {date.today().strftime("%d %B %Y")} if no date was specified

{("The firm's knowledge base contains the following relevant reference:\n" + firm_knowledge) if firm_knowledge else ""}

Output the complete contract text only. No preamble or commentary."""

    resp = get_anthropic().messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=3000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}]
    )
    return resp.content[0].text


def export_contract_docx(contract_text: str, title: str) -> bytes:
    """Convert contract text to a properly formatted DOCX."""
    if not DOCX_OK:
        return None
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Body — parse line by line
    for line in contract_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        para = doc.add_paragraph()
        # Detect review flags
        if "[REVIEW]" in line:
            line = line.replace("[REVIEW]", "")
            run = para.add_run("[REVIEW] " + line)
            run.font.color.rgb = RGBColor(0x8B, 0x26, 0x35)
            run.bold = True
        # Numbered clause heading
        elif re.match(r"^\d+\.\s+[A-Z]", line):
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(10.5)
        else:
            run = para.add_run(line)
            run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = Pt(14)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_contract_pdf(contract_text: str, title: str) -> bytes:
    """Convert contract text to PDF using reportlab."""
    if not REPORTLAB_OK:
        return None
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
          leftMargin=2.5*cm, rightMargin=2.5*cm,
          topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"],
        fontSize=13, spaceAfter=12, alignment=1)
    body_style  = ParagraphStyle("B", parent=styles["Normal"],
        fontSize=9.5, leading=14, spaceAfter=4)
    flag_style  = ParagraphStyle("F", parent=styles["Normal"],
        fontSize=9.5, leading=14, spaceAfter=4, textColor=(0.55,0.15,0.21))

    story = [Paragraph(title.upper(), title_style), Spacer(1, 0.3*cm)]
    for line in contract_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.15*cm))
        elif "[REVIEW]" in line:
            story.append(Paragraph("[REVIEW] " + line.replace("[REVIEW]","").strip(), flag_style))
        else:
            story.append(Paragraph(line, body_style))
    doc.build(story)
    return buf.getvalue()


# ---- RAG Q&A -----------------------------------------------------------
def rag_ask(question: str, case_id: str = None, history: list = None) -> tuple:
    c = ctx()
    vector = get_voyage().embed([question], model="voyage-law-2", input_type="query").embeddings[0]
    params = {"query_embedding": vector, "match_count": 5}
    if case_id:
        params["p_case_id"] = case_id

    chunks = ScopedDB(c["firm_id"]).rpc("match_documents", params).execute().data or []

    if chunks:
        lines = ["--- RETRIEVED CONTEXT ---\n"]
        for i, ch in enumerate(chunks, 1):
            src = ch.get("metadata", {}).get("source", "unknown")
            lines.append(f"[{i}] {src}\n{ch.get('content','')}\n")
        lines.append("--- END CONTEXT ---")
        context = "\n".join(lines)
    else:
        context = "No relevant documents found in the knowledge base for this query."

    msgs = (history or []) + [{"role": "user", "content": question}]
    resp = get_anthropic().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1500,
        system=[{"type": "text", "text": SYSTEM_PROMPT + "\n\n" + context,
                 "cache_control": {"type": "ephemeral"}}],
        messages=msgs)
    answer = resp.content[0].text

    try:
        sdb().table("conversations").insert({
            "user_id": c["user_id"], "case_id": case_id,
            "question": question, "answer": answer,
            "doc_sources": [ch.get("metadata", {}).get("source") for ch in chunks],
        }).execute()
    except Exception:
        pass

    return answer, len(chunks)

def call_claude(prompt: str, history: list = None) -> str:
    msgs = (history or []) + [{"role": "user", "content": prompt}]
    resp = get_anthropic().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1500,
        system=[{"type": "text", "text": SYSTEM_PROMPT, "cache_control": {"type": "ephemeral"}}],
        messages=msgs)
    return resp.content[0].text

def pdf_page_images(file_bytes: bytes, max_pages: int = 3) -> list:
    if not PYMUPDF_OK: return []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    mat = fitz.Matrix(120/72, 120/72)
    imgs = []
    for i, page in enumerate(doc):
        if i >= max_pages: break
        pix = page.get_pixmap(matrix=mat)
        imgs.append(base64.standard_b64encode(pix.tobytes("png")).decode())
    doc.close()
    return imgs


# ---- Login ------------------------------------------------------------
def show_login():
    st.markdown("""<div class="login-card">
        <div class="wordmark">AltoLex</div>
        <div class="tagline">by Altovian</div>
        </div>""", unsafe_allow_html=True)
    with st.form("login"):
        st.subheader("Sign in")
        email    = st.text_input("Email address")
        password = st.text_input("Password", type="password")
        if st.form_submit_button("Sign in ->", use_container_width=True):
            result = login(email, password)
            if result:
                st.session_state["ctx"] = result; st.rerun()
            else:
                st.error("Invalid email or password.")
    st.markdown('<div class="disclaimer" style="max-width:420px;margin:0 auto">AltoLex provides legal information only — not legal advice. All outputs must be reviewed by a qualified attorney.</div>', unsafe_allow_html=True)


# ---- Sidebar ----------------------------------------------------------
def show_sidebar():
    c = ctx()
    with st.sidebar:
        st.markdown('<div style="font-family:\'Playfair Display\',serif;font-size:22px;font-weight:600;color:#d4aa7a">AltoLex</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:10px;text-transform:uppercase;letter-spacing:0.1em;color:rgba(232,228,220,0.4);margin-bottom:16px">by Altovian</div>', unsafe_allow_html=True)
        st.markdown(f"**{c.get('full_name','')}**")
        st.markdown(f"<small>{c.get('email','')}</small>", unsafe_allow_html=True)
        bc = "badge-admin" if c.get("role") == "admin" else "badge-firm"
        st.markdown(f'<span class="badge {bc}">{c.get("role","").upper()}</span>', unsafe_allow_html=True)
        st.divider()

        modules = ["📋 Client Intake", "👥 Clients & Cases",
                   "📁 Document Library", "💬 AI Search", "⚖️ Common Knowledge",
                   "✍️ Draft Contract"]
        if is_admin():
            modules.append("⚙ Admin")

        module = st.radio("Module", modules, label_visibility="collapsed")
        st.divider()
        st.markdown("**System**")
        st.markdown(f"{'OK' if PYPDF2_OK else 'X'} Text PDFs")
        st.markdown(f"{'OK' if PYMUPDF_OK else 'X'} Scanned PDFs (OCR)")
        st.markdown(f"{'OK' if DOCX_OK else 'X'} Word docs")
        st.markdown(f"{'OK' if VOYAGE_OK else 'X'} voyage-law-2")
        st.markdown(f"{'OK' if REPORTLAB_OK else 'X'} PDF export")
        st.divider()
        if st.button("Sign out"):
            logout(c.get("token")); st.session_state.clear(); st.rerun()
        st.markdown('<div class="disclaimer">Legal information only. Attorney review required on all outputs.</div>', unsafe_allow_html=True)
    return module


# ---- Admin panel -------------------------------------------------------
def show_admin():
    c = ctx()
    st.title("Admin - User Management")
    tab_users, tab_add = st.tabs(["Users", "Add User"])

    with tab_users:
        users = list_users(c["firm_id"])
        st.markdown(f"**{len(users)} users**")
        st.divider()
        for u in users:
            with st.expander(f"{'Active' if u['is_active'] else 'Inactive'} | {u['full_name']} - {u['email']} ({u['role']})"):
                col1, col2, col3 = st.columns([2, 2, 1])
                with col1:
                    new_name = st.text_input("Full name", value=u["full_name"], key=f"n_{u['id']}")
                    new_role = st.selectbox("Role", ROLES, index=ROLES.index(u["role"]), key=f"r_{u['id']}")
                with col2:
                    new_pw = st.text_input("New password (blank = no change)", type="password", key=f"p_{u['id']}")
                    st.caption(f"Last login: {u.get('last_login', 'Never')}")
                with col3:
                    st.markdown("&nbsp;", unsafe_allow_html=True)
                    if st.button("Save", key=f"s_{u['id']}"):
                        upd = {"full_name": new_name, "role": new_role}
                        if new_pw.strip(): upd["password"] = new_pw.strip()
                        update_user(c["firm_id"], u["id"], upd)
                        st.success("Saved."); st.rerun()
                    if u["id"] != c["user_id"]:
                        if u["is_active"]:
                            if st.button("Deactivate", key=f"d_{u['id']}"):
                                deactivate_user(c["firm_id"], u["id"]); st.rerun()
                        else:
                            if st.button("Reactivate", key=f"d_{u['id']}"):
                                reactivate_user(c["firm_id"], u["id"]); st.rerun()

    with tab_add:
        st.subheader("Add a new user")
        with st.form("add_user"):
            col1, col2 = st.columns(2)
            with col1:
                nfn = st.text_input("Full name *")
                nem = st.text_input("Email *")
            with col2:
                nrl = st.selectbox("Role *", ROLES, index=2)
                npw = st.text_input("Temporary password *", type="password")
            cpw = st.text_input("Confirm password *", type="password")
            if st.form_submit_button("Create user ->", use_container_width=True):
                errs = []
                if not nfn.strip(): errs.append("Full name required.")
                if not nem.strip(): errs.append("Email required.")
                if not npw.strip(): errs.append("Password required.")
                if npw != cpw:      errs.append("Passwords do not match.")
                if len(npw) < 8:    errs.append("Password must be at least 8 characters.")
                if errs:
                    for e in errs: st.error(e)
                else:
                    try:
                        create_user(c["firm_id"], nem.strip(), npw, nfn.strip(), nrl, c["user_id"])
                        log_action(c["firm_id"], c["user_id"], "user_create", {"email": nem})
                        st.success(f"User {nfn} created as {nrl}.")
                    except ValueError as e:
                        st.error(str(e))


# ======================================================================
# SESSION GATE
# ======================================================================
for k, v in [("ctx", {}), ("chat_history", []), ("_last_scope", "")]:
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state["ctx"]:
    fresh = validate_session(st.session_state["ctx"].get("token", ""))
    if fresh:
        st.session_state["ctx"] = fresh
    else:
        st.session_state.clear(); st.rerun()

if not st.session_state["ctx"]:
    show_login(); st.stop()

c      = ctx()
module = show_sidebar()


# ======================================================================
# MODULE: CLIENT INTAKE
# ======================================================================
if module == "📋 Client Intake":
    if is_readonly():
        st.warning("Readonly role cannot submit intake forms."); st.stop()

    st.title("New Client Intake")
    col1, col2 = st.columns(2)
    with col1:
        name    = st.text_input("Full name *")
        email_i = st.text_input("Email")
        phone_i = st.text_input("Phone", placeholder="+94 77 xxx xxxx")
    with col2:
        area    = st.selectbox("Area of law", [""] + AREAS_OF_LAW)
        urgency = st.selectbox("Urgency", ["", "Urgent - within 48 hours", "Within the week", "No immediate deadline"])
        prior   = st.selectbox("Prior legal action?", ["", "No, first time", "Yes - previous attorney", "Yes - in proceedings"])

    matter = st.text_area("Describe the matter *", height=120)
    save_client = st.checkbox("Save as new client record in the system", value=True)

    if st.button("Analyse & Triage ->", type="primary"):
        if not name or not area or not matter:
            st.warning("Please fill in name, area of law, and matter description.")
        else:
            with st.spinner("Analysing..."):
                prompt = (f"Triage this new client intake:\n\nClient: {name}\nArea: {area}\n"
                          f"Urgency: {urgency or 'Not specified'}\nPrior: {prior or 'Not specified'}\n"
                          f"Matter:\n\"{matter}\"\n\nProvide a full intake triage report.")
                result = call_claude(prompt)
                if save_client:
                    cl = create_client_record(name, email_i, phone_i)
                    create_case_record(cl["id"], f"{area} - {name}", area)
                    log_action(c["firm_id"], c["user_id"], "intake", {"client_id": cl["id"]})
                    st.success(f"Client **{name}** and case saved. Go to Document Library to upload files.")
            st.divider()
            st.subheader("Intake Analysis")
            st.caption("AI generated - attorney review required")
            st.markdown(result)


# ======================================================================
# MODULE: CLIENTS & CASES
# ======================================================================
elif module == "👥 Clients & Cases":
    st.title("Clients & Cases")
    tab_list, tab_new = st.tabs(["All Clients", "New Client"])

    with tab_list:
        clients = get_clients()
        if not clients:
            st.info("No clients yet. Use the New Client tab or create one during intake.")
        else:
            for cl in clients:
                with st.expander(f"{cl['full_name']}  {'| ' + cl['email'] if cl.get('email') else ''}"):
                    st.markdown(f"**Email:** {cl.get('email') or '-'}  |  **Phone:** {cl.get('phone') or '-'}")
                    cases = get_cases(cl["id"])
                    if cases:
                        st.markdown(f"**{len(cases)} case(s):**")
                        for ca in cases:
                            icon = "Open" if ca["status"] == "open" else "Closed"
                            st.markdown(f"- [{icon}] **{ca['title']}** · {ca['area_of_law']}")
                    else:
                        st.markdown("*No cases yet.*")

                    with st.form(f"addcase_{cl['id']}"):
                        st.markdown("**Add a case:**")
                        ct = st.text_input("Case title", key=f"ct_{cl['id']}")
                        ca_area = st.selectbox("Area of law", AREAS_OF_LAW, key=f"caa_{cl['id']}")
                        if st.form_submit_button("Add case"):
                            if ct.strip():
                                create_case_record(cl["id"], ct.strip(), ca_area)
                                st.success("Case added."); st.rerun()

    with tab_new:
        st.subheader("New client")
        with st.form("new_client"):
            col1, col2 = st.columns(2)
            with col1:
                nc_name  = st.text_input("Full name *")
                nc_email = st.text_input("Email")
            with col2:
                nc_phone = st.text_input("Phone")
                nc_area  = st.selectbox("Initial area of law", [""] + AREAS_OF_LAW)
            nc_case = st.text_input("Initial case title (optional)")
            if st.form_submit_button("Create client ->", use_container_width=True):
                if not nc_name.strip():
                    st.error("Full name is required.")
                else:
                    cl = create_client_record(nc_name.strip(), nc_email, nc_phone)
                    if nc_case.strip() and nc_area:
                        create_case_record(cl["id"], nc_case.strip(), nc_area)
                    log_action(c["firm_id"], c["user_id"], "client_create", {"name": nc_name})
                    st.success(f"Client **{nc_name}** created."); st.rerun()


# ======================================================================
# MODULE: DOCUMENT LIBRARY
# ======================================================================
elif module == "📁 Document Library":
    st.title("Document Library")

    clients = get_clients()
    if not clients:
        st.warning("No clients yet. Create a client first in **Clients & Cases**.")
        st.stop()

    col1, col2 = st.columns(2)
    with col1:
        client_map = {cl["full_name"]: cl for cl in clients}
        sel_cl = client_map[st.selectbox("Client", list(client_map.keys()))]
    with col2:
        cases = get_cases(sel_cl["id"])
        if not cases:
            st.warning("This client has no cases. Add one in Clients & Cases first.")
            st.stop()
        case_map = {ca["title"]: ca for ca in cases}
        sel_ca = case_map[st.selectbox("Case", list(case_map.keys()))]

    tab_docs, tab_upload = st.tabs(["📂 Documents", "⬆ Upload"])

    with tab_docs:
        existing = get_stored_docs_with_path(case_id=sel_ca["id"])
        if not existing:
            st.info("No documents stored for this case yet. Use the Upload tab.")
        else:
            st.markdown(f"**{len(existing)} document(s)** for *{sel_ca['title']}*")
            st.divider()

            # Preview selector
            doc_names   = [d["name"] for d in existing]
            preview_sel = st.selectbox("Select document to preview", doc_names,
                                       key="doc_preview_sel")
            sel_doc     = next((d for d in existing if d["name"] == preview_sel), None)

            if sel_doc:
                col_meta, col_dl = st.columns([3, 1])
                with col_meta:
                    dtype = sel_doc.get("doc_type", "general")
                    badge = "badge-common" if dtype == "ordinance" else "badge-case"
                    st.markdown(
                        f'<span class="badge {badge}">{dtype}</span> '
                        f'<small>uploaded {sel_doc["created_at"][:10]}</small>',
                        unsafe_allow_html=True
                    )

                # Load bytes from Storage
                if sel_doc.get("storage_path"):
                    with st.spinner("Loading document..."):
                        doc_bytes = storage_download(sel_doc["storage_path"])
                    if doc_bytes:
                        with col_dl:
                            st.download_button(
                                "⬇ Download",
                                data=doc_bytes,
                                file_name=sel_doc["name"],
                                mime="application/octet-stream",
                                use_container_width=True,
                            )
                        st.divider()
                        if sel_doc["name"].lower().endswith(".pdf"):
                            render_pdf_viewer(doc_bytes, sel_doc["name"], height=580)
                        else:
                            # DOCX / other — show extracted text
                            text, _ = extract_text(doc_bytes, sel_doc["name"])
                            st.text_area("Document content", value=text[:5000] +
                                ("..." if len(text) > 5000 else ""),
                                height=480, disabled=True, label_visibility="collapsed")
                    else:
                        st.warning("Document file not found in storage. "
                                   "Re-upload to enable preview.")
                else:
                    st.info("Preview not available — document was uploaded before "
                            "storage was enabled. Re-upload to enable preview.")

    with tab_upload:
        st.markdown("Files are embedded for AI search and stored for preview.")
        uploaded_files = st.file_uploader(
            "Choose files", type=["pdf", "docx", "doc"],
            accept_multiple_files=True, label_visibility="collapsed"
        )
        if uploaded_files:
            label = f"Embed & Store {len(uploaded_files)} file(s)"
            if st.button(label, type="primary", disabled=is_readonly()):
                for uf in uploaded_files:
                    with st.spinner(f"Processing {uf.name}..."):
                        file_bytes = uf.read()
                        if uf.name.lower().endswith(".pdf") and PYMUPDF_OK:
                            imgs = pdf_page_images(file_bytes, max_pages=1)
                            if imgs:
                                st.image(base64.b64decode(imgs[0]),
                                         width=240, caption=uf.name)
                        result = embed_and_store(file_bytes, uf.name,
                                                 c["firm_id"], sel_ca["id"])
                    if result.get("skipped"):
                        st.warning(f"Skipped — {uf.name} already stored.")
                    elif result.get("chunks", 0) > 0:
                        st.success(f"Stored {uf.name} — {result['chunks']} chunks "
                                   f"({result['method']})")
                        log_action(c["firm_id"], c["user_id"], "ingest",
                                   {"filename": uf.name, "case_id": sel_ca["id"]})
                    else:
                        st.error(f"Failed {uf.name} — {result.get('error','')}")
                st.rerun()


# ======================================================================
# MODULE: AI SEARCH (RAG)
# ======================================================================
elif module == "💬 AI Search":
    st.title("AI Search")
    st.markdown("Ask questions grounded in stored documents. Searches both case documents and the common knowledge base.")

    scope = st.radio("Search scope", ["Specific case", "All firm documents"], horizontal=True)

    sel_case_id = None
    sel_ca_title = "all documents"
    if scope == "Specific case":
        clients = get_clients()
        if not clients:
            st.info("No clients yet."); st.stop()
        col1, col2 = st.columns(2)
        with col1:
            client_map = {cl["full_name"]: cl for cl in clients}
            sel_cl = client_map[st.selectbox("Client", list(client_map.keys()))]
        with col2:
            cases = get_cases(sel_cl["id"])
            if not cases:
                st.warning("No cases for this client."); st.stop()
            case_map = {ca["title"]: ca for ca in cases}
            sel_ca = case_map[st.selectbox("Case", list(case_map.keys()))]
            sel_case_id  = sel_ca["id"]
            sel_ca_title = sel_ca["title"]

    scope_key = f"{scope}_{sel_case_id}"
    if st.session_state.get("_last_scope") != scope_key:
        st.session_state.chat_history = []
        st.session_state["_last_scope"] = scope_key

    st.divider()

    if not st.session_state.chat_history:
        with st.chat_message("assistant", avatar="⚖️"):
            st.markdown(f"Ready to search **{sel_ca_title}**. What would you like to know?")

    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"], avatar="⚖️" if msg["role"] == "assistant" else "👤"):
            st.markdown(msg["content"])
            if msg.get("chunks_used") is not None:
                st.caption(f"{msg['chunks_used']} document chunks retrieved")

    if user_input := st.chat_input("Ask about these documents..."):
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user", avatar="👤"):
            st.markdown(user_input)
        with st.chat_message("assistant", avatar="⚖️"):
            with st.spinner("Searching documents..."):
                history = [{"role": m["role"], "content": m["content"]}
                           for m in st.session_state.chat_history[:-1][-8:]]
                answer, n = rag_ask(user_input, sel_case_id, history)
            st.markdown(answer)
            if n > 0:
                st.caption(f"{n} document chunks retrieved")
                # Offer source preview inline
                if "source_chunks" not in st.session_state:
                    st.session_state["source_chunks"] = []
            st.session_state.chat_history.append({
                "role": "assistant", "content": answer, "chunks_used": n
            })

    if st.session_state.chat_history:
        if st.button("Clear conversation"):
            st.session_state.chat_history = []; st.rerun()


# ======================================================================
# MODULE: COMMON KNOWLEDGE (Ordinances / Firm-wide)
# ======================================================================
elif module == "⚖️ Common Knowledge":
    st.title("Common Knowledge Base")
    st.markdown("Firm-wide documents — **ordinances, acts, regulations, standard templates** — available for any client search. Not tied to any specific case.")

    tab_up, tab_browse = st.tabs(["Upload", "Browse"])

    with tab_up:
        if is_readonly():
            st.warning("Readonly role cannot upload documents.")
        else:
            common_files = st.file_uploader(
                "Upload ordinances, acts, or firm templates",
                type=["pdf", "docx", "doc"],
                accept_multiple_files=True,
                label_visibility="collapsed"
            )
            if common_files:
                if st.button(f"Store {len(common_files)} file(s) as Common Knowledge", type="primary"):
                    for uf in common_files:
                        with st.spinner(f"Processing {uf.name}..."):
                            file_bytes = uf.read()
                            if uf.name.lower().endswith(".pdf") and PYMUPDF_OK:
                                imgs = pdf_page_images(file_bytes, max_pages=1)
                                if imgs:
                                    st.image(base64.b64decode(imgs[0]), width=260, caption=uf.name)
                            result = embed_and_store(file_bytes, uf.name, c["firm_id"], case_id=None)
                        if result.get("skipped"):
                            st.warning(f"Skipped {uf.name} - already stored.")
                        elif result.get("chunks", 0) > 0:
                            st.success(f"Stored {uf.name} - {result['chunks']} chunks")
                            log_action(c["firm_id"], c["user_id"], "ingest_common", {"filename": uf.name})
                        else:
                            st.error(f"Failed {uf.name} - {result.get('error','could not extract text')}")
                    st.rerun()

    with tab_browse:
        common_docs = get_stored_docs(firm_wide=True)
        if not common_docs:
            st.info("No common documents yet. Upload ordinances and templates in the Upload tab.")
        else:
            st.markdown(f"**{len(common_docs)} common document(s):**")
            for d in common_docs:
                dtype = doc_type_tag(d["name"])
                badge = "badge-common" if dtype in ("ordinance", "sop") else "badge-case"
                st.markdown(
                    f'<span class="badge {badge}">{dtype}</span> {d["name"]} '
                    f'<small>uploaded {d["created_at"][:10]}</small>',
                    unsafe_allow_html=True
                )


# ======================================================================
# MODULE: DRAFT CONTRACT
# ======================================================================
elif module == "✍️ Draft Contract":
    if is_readonly():
        st.warning("Readonly role cannot draft contracts."); st.stop()

    st.title("✍️ Draft Contract")
    st.markdown("Select a template, fill in the details, and let AltoLex generate a complete contract ready for attorney review.")

    # Template selector
    tmpl_name = st.selectbox(
        "Contract type",
        list(CONTRACT_TEMPLATES.keys()),
        key="tmpl_sel"
    )
    tmpl = CONTRACT_TEMPLATES[tmpl_name]
    st.caption(tmpl["description"])
    st.divider()

    # Optional: link to a client/case
    clients = get_clients()
    link_col1, link_col2, link_col3 = st.columns([2, 2, 1])
    with link_col1:
        link_client = st.selectbox("Link to client (optional)",
            ["— none —"] + [cl["full_name"] for cl in clients], key="draft_cl")
    sel_draft_case = None
    with link_col2:
        if link_client != "— none —":
            cl_obj = next((x for x in clients if x["full_name"] == link_client), None)
            if cl_obj:
                cl_cases = get_cases(cl_obj["id"])
                if cl_cases:
                    link_case = st.selectbox("Case", [ca["title"] for ca in cl_cases],
                                             key="draft_ca")
                    sel_draft_case = next((ca for ca in cl_cases
                                           if ca["title"] == link_case), None)

    st.divider()

    # Variable fields
    st.subheader("Contract details")
    variables = {}
    # Render fields in 2-column grid
    field_pairs = [tmpl["fields"][i:i+2] for i in range(0, len(tmpl["fields"]), 2)]
    for pair in field_pairs:
        cols = st.columns(len(pair))
        for col, field in zip(cols, pair):
            with col:
                key      = field["key"]
                label    = field["label"]
                ftype    = field.get("type", "text")
                default  = field.get("default", "")
                placeholder = field.get("placeholder", "")
                if ftype == "text":
                    variables[key] = st.text_input(label, value=str(default),
                                                   placeholder=placeholder,
                                                   key=f"f_{tmpl_name}_{key}")
                elif ftype == "number":
                    variables[key] = st.number_input(label, value=int(default) if default else 1,
                                                     min_value=0, key=f"f_{tmpl_name}_{key}")
                elif ftype == "date":
                    variables[key] = str(st.date_input(label,
                                         value=date.today(),
                                         key=f"f_{tmpl_name}_{key}"))
                elif ftype == "textarea":
                    variables[key] = st.text_area(label, value=str(default),
                                                  placeholder=placeholder, height=80,
                                                  key=f"f_{tmpl_name}_{key}")

    st.divider()

    # Draft button
    if "drafted_contract" not in st.session_state:
        st.session_state["drafted_contract"] = ""
    if "drafted_title"    not in st.session_state:
        st.session_state["drafted_title"]    = ""

    if st.button("✦ Generate Contract Draft", type="primary"):
        filled = {k: v for k, v in variables.items() if str(v).strip()}
        if len(filled) < 2:
            st.warning("Please fill in at least the main party names before drafting.")
        else:
            with st.spinner("Drafting contract..."):
                draft = draft_contract_with_ai(tmpl_name, filled)
                st.session_state["drafted_contract"] = draft
                st.session_state["drafted_title"]    = tmpl_name
                log_action(c["firm_id"], c["user_id"], "contract_draft",
                           {"template": tmpl_name,
                            "case_id": sel_draft_case["id"] if sel_draft_case else None})
            st.success("Draft generated. Review carefully before use.")

    # Show draft
    if st.session_state.get("drafted_contract"):
        draft_text  = st.session_state["drafted_contract"]
        draft_title = st.session_state["drafted_title"]

        tab_preview, tab_edit, tab_export = st.tabs(["👁 Preview", "✏️ Edit", "⬇ Export"])

        with tab_preview:
            st.caption("Clauses marked [REVIEW] require specific attorney attention.")
            # Colour-code [REVIEW] flags
            lines = draft_text.split("\n")
            for line in lines:
                if "[REVIEW]" in line:
                    st.markdown(
                        f'<div style="background:rgba(139,38,53,0.06);border-left:3px solid '
                        f'rgba(139,38,53,0.4);padding:6px 10px;border-radius:0 6px 6px 0;'
                        f'margin:2px 0;font-size:13px;color:var(--text-color)">'
                        f'⚠️ {line.replace("[REVIEW]","").strip()}</div>',
                        unsafe_allow_html=True
                    )
                elif line.strip():
                    st.markdown(line)
                else:
                    st.markdown("")

        with tab_edit:
            st.caption("Edit the draft directly. Changes are saved automatically.")
            edited = st.text_area("Contract text", value=draft_text,
                                  height=600, label_visibility="collapsed",
                                  key="contract_edit_area")
            if edited != draft_text:
                st.session_state["drafted_contract"] = edited
                draft_text = edited

        with tab_export:
            st.markdown("Download the contract in your preferred format.")
            st.caption("⚠️ Always review the AI-generated draft before sending to any party.")
            ecol1, ecol2 = st.columns(2)
            with ecol1:
                docx_bytes = export_contract_docx(draft_text, draft_title)
                if docx_bytes:
                    st.download_button(
                        "⬇ Download as Word (.docx)",
                        data=docx_bytes,
                        file_name=f"{draft_title.replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                else:
                    st.warning("python-docx not installed — DOCX export unavailable.")
            with ecol2:
                pdf_bytes = export_contract_pdf(draft_text, draft_title)
                if pdf_bytes:
                    st.download_button(
                        "⬇ Download as PDF",
                        data=pdf_bytes,
                        file_name=f"{draft_title.replace(' ','_')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                else:
                    st.warning("reportlab not installed — PDF export unavailable.")

            # Optionally save to case documents
            if sel_draft_case and (docx_bytes or pdf_bytes):
                st.divider()
                if st.button("💾 Save draft to case documents"):
                    save_bytes = docx_bytes or pdf_bytes
                    save_name  = f"{draft_title.replace(' ','_')}_DRAFT.{'docx' if docx_bytes else 'pdf'}"
                    result = embed_and_store(save_bytes, save_name,
                                            c["firm_id"], sel_draft_case["id"])
                    if result.get("skipped"):
                        st.info("Draft already saved to this case.")
                    elif result.get("chunks", 0) > 0:
                        st.success(f"Draft saved to case '{sel_draft_case['title']}' "
                                   f"and indexed for AI search.")
                    else:
                        st.error("Could not save draft to case.")


# ======================================================================
# MODULE: ADMIN
# ======================================================================
elif module == "⚙ Admin":
    if not is_admin():
        st.error("Access denied."); st.stop()
    show_admin()
