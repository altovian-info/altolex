"""
AltoLex — app_v3.py
Custom authentication — users/roles in your own Supabase tables.
No Supabase Auth. No auth.users dependency.
Includes built-in admin panel for managing users and firms.

Run: streamlit run app_v3.py
"""

import streamlit as st
import anthropic
import os, io, base64
from pathlib import Path
from auth import (login, logout, validate_session,
                  create_user, update_user, deactivate_user, list_users,
                  log_action)

try:
    import fitz;        PYMUPDF_OK = True
except ImportError:     PYMUPDF_OK = False
try:
    from docx import Document as DocxDocument; DOCX_OK = True
except ImportError:     DOCX_OK = False
try:
    import PyPDF2;      PYPDF2_OK = True
except ImportError:     PYPDF2_OK = False


st.set_page_config(page_title="AltoLex", page_icon="⚖️", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600&family=DM+Sans:wght@400;500&display=swap');
html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
h1, h2, h3 { font-family: 'Playfair Display', serif !important; }
.stApp { background: #f7f4ef; }
.badge { display:inline-block; padding:3px 10px; border-radius:20px; font-size:11px;
         font-weight:600; letter-spacing:0.06em; text-transform:uppercase; }
.badge-firm  { background:rgba(184,147,90,0.12); color:#b8935a; border:1px solid rgba(184,147,90,0.3); }
.badge-case  { background:rgba(42,90,180,0.1);   color:#2a5ab4; border:1px solid rgba(42,90,180,0.2); }
.badge-admin { background:rgba(139,38,53,0.08);  color:#8b2635; border:1px solid rgba(139,38,53,0.2); }
.badge-ok    { background:rgba(45,106,79,0.1);   color:#2d6a4f; border:1px solid rgba(45,106,79,0.2); }
.disclaimer  { background:rgba(184,147,90,0.08); border:1px solid rgba(184,147,90,0.2);
               border-radius:8px; padding:12px 16px; font-size:12px; color:#6b6b80; line-height:1.6; }
.user-row    { background:white; border:1px solid rgba(184,147,90,0.15); border-radius:8px;
               padding:12px 16px; margin-bottom:8px; }
</style>
""", unsafe_allow_html=True)

SYSTEM_PROMPT = """You are AltoLex, a professional legal information assistant.
You have access to the firm's document library — retrieved context is provided below.
Prefer information from the context. Cite source filenames when referencing them.
RULES:
1. Never give definitive legal advice
2. Always recommend attorney review
3. If context lacks relevant information, say so — do not fabricate
4. Flag urgent deadlines prominently
5. End substantive responses with: "This is legal information only. Please consult a qualified attorney."
"""

ROLES = ["admin", "partner", "associate", "paralegal", "readonly"]


# ── Anthropic client ──────────────────────────────────────────────────────────
@st.cache_resource
def get_anthropic():
    key = os.environ.get("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY","")
    if not key: st.error("ANTHROPIC_API_KEY not set."); st.stop()
    return anthropic.Anthropic(api_key=key)


# ── Session helpers ───────────────────────────────────────────────────────────
def ctx() -> dict:
    return st.session_state.get("ctx", {})

def is_admin() -> bool:
    return ctx().get("role") == "admin"

def is_readonly() -> bool:
    return ctx().get("role") == "readonly"


# ── Login screen ──────────────────────────────────────────────────────────────
def show_login():
    st.markdown("""
    <div style="max-width:420px;margin:60px auto 24px;background:white;
                border:1px solid rgba(184,147,90,0.2);border-radius:12px;padding:32px">
        <div style="font-family:'Playfair Display',serif;font-size:28px;
                    font-weight:600;color:#1a1a2e;margin-bottom:3px">AltoLex</div>
        <div style="font-size:10px;text-transform:uppercase;letter-spacing:0.1em;
                    color:#6b6b80;margin-bottom:28px">by Altovian</div>
    </div>
    """, unsafe_allow_html=True)

    with st.form("login"):
        st.subheader("Sign in")
        email    = st.text_input("Email address")
        password = st.text_input("Password", type="password")
        if st.form_submit_button("Sign in →", use_container_width=True):
            result = login(email, password)
            if result:
                st.session_state["ctx"] = result
                st.rerun()
            else:
                st.error("Invalid email or password.")

    st.markdown('<div class="disclaimer" style="max-width:420px;margin:0 auto">⚠ AltoLex provides legal information only — not legal advice. All outputs must be reviewed by a qualified attorney.</div>', unsafe_allow_html=True)


# ── Sidebar ───────────────────────────────────────────────────────────────────
def show_sidebar():
    c = ctx()
    with st.sidebar:
        st.markdown('<div style="font-family:\'Playfair Display\',serif;font-size:22px;font-weight:600;color:#d4aa7a">AltoLex</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:10px;text-transform:uppercase;letter-spacing:0.1em;color:rgba(232,228,220,0.4);margin-bottom:16px">by Altovian</div>', unsafe_allow_html=True)
        st.markdown(f"**{c.get('full_name','')}**")
        st.markdown(f"<small style='color:#6b6b80'>{c.get('email','')}</small>", unsafe_allow_html=True)
        badge_class = "badge-admin" if c.get("role")=="admin" else "badge-firm"
        st.markdown(f'<span class="badge {badge_class}">{c.get("role","").upper()}</span>', unsafe_allow_html=True)
        st.divider()

        modules = ["📋 Client Intake", "💬 Legal Q&A", "📄 Document Review"]
        if is_admin():
            modules.append("⚙ Admin")

        module = st.radio("Module", modules, label_visibility="collapsed")

        st.divider()
        st.markdown("**System**")
        st.markdown(f"{'✅' if PYPDF2_OK else '❌'} Text PDFs")
        st.markdown(f"{'✅' if PYMUPDF_OK else '❌'} Scanned PDFs")
        st.markdown(f"{'✅' if DOCX_OK else '❌'} Word docs")
        st.divider()

        if st.button("Sign out"):
            logout(c.get("token"))
            st.session_state.clear()
            st.rerun()

        st.markdown('<div class="disclaimer">⚠ Legal information only. Attorney review required on all outputs.</div>', unsafe_allow_html=True)

    return module


# ── Admin panel ───────────────────────────────────────────────────────────────
def show_admin():
    c = ctx()
    st.title("⚙ Admin — User Management")
    st.markdown(f"Managing users for your firm. Only **admin** role can access this panel.")

    tab_users, tab_add = st.tabs(["👥 Users", "➕ Add User"])

    # ── Tab: existing users ──
    with tab_users:
        users = list_users(c["firm_id"])
        if not users:
            st.info("No users yet. Use the 'Add User' tab to create the first user.")
        else:
            st.markdown(f"**{len(users)} user{'s' if len(users)!=1 else ''}** in your firm")
            st.divider()
            for u in users:
                with st.expander(f"{'🟢' if u['is_active'] else '⚫'} {u['full_name']}  —  {u['email']}  ·  {u['role'].upper()}"):
                    col1, col2, col3 = st.columns([2,2,1])
                    with col1:
                        new_name = st.text_input("Full name", value=u["full_name"], key=f"name_{u['id']}")
                        new_role = st.selectbox("Role", ROLES, index=ROLES.index(u["role"]), key=f"role_{u['id']}")
                    with col2:
                        new_pw = st.text_input("New password (leave blank to keep)", type="password", key=f"pw_{u['id']}")
                        st.markdown(f"<small style='color:#6b6b80'>Last login: {u.get('last_login','Never')}</small>", unsafe_allow_html=True)
                    with col3:
                        st.markdown("&nbsp;", unsafe_allow_html=True)
                        if st.button("💾 Save", key=f"save_{u['id']}"):
                            # Prevent removing the last admin
                            if u["role"] == "admin" and new_role != "admin":
                                admins = [x for x in users if x["role"]=="admin" and x["is_active"]]
                                if len(admins) <= 1:
                                    st.error("Cannot remove the last admin.")
                                    st.stop()
                            updates = {"full_name": new_name, "role": new_role}
                            if new_pw.strip(): updates["password"] = new_pw.strip()
                            update_user(u["id"], updates)
                            log_action(c["firm_id"], c["user_id"], "user_update",
                                       {"target_user": u["id"], "new_role": new_role})
                            st.success("Saved.")
                            st.rerun()

                        # Deactivate / reactivate
                        if u["id"] != c["user_id"]:  # can't deactivate yourself
                            label  = "🔴 Deactivate" if u["is_active"] else "🟢 Reactivate"
                            if st.button(label, key=f"deact_{u['id']}"):
                                if u["is_active"]:
                                    admins = [x for x in users if x["role"]=="admin" and x["is_active"]]
                                    if u["role"]=="admin" and len(admins)<=1:
                                        st.error("Cannot deactivate the last admin.")
                                        st.stop()
                                    deactivate_user(u["id"])
                                    log_action(c["firm_id"], c["user_id"], "user_deactivate",
                                               {"target_user": u["id"]})
                                else:
                                    from auth import _svc
                                    _svc().table("users").update({"is_active":True}).eq("id",u["id"]).execute()
                                    log_action(c["firm_id"], c["user_id"], "user_reactivate",
                                               {"target_user": u["id"]})
                                st.rerun()

    # ── Tab: add user ──
    with tab_add:
        st.subheader("Add a new user")
        with st.form("add_user"):
            col1, col2 = st.columns(2)
            with col1:
                new_full_name = st.text_input("Full name *")
                new_email     = st.text_input("Email address *")
            with col2:
                new_role_sel  = st.selectbox("Role *", ROLES, index=2)  # default: associate
                new_password  = st.text_input("Temporary password *", type="password",
                                              help="User should change this on first login")
            confirm_pw = st.text_input("Confirm password *", type="password")

            if st.form_submit_button("Create user →", use_container_width=True):
                errors = []
                if not new_full_name.strip(): errors.append("Full name is required.")
                if not new_email.strip():     errors.append("Email is required.")
                if not new_password.strip():  errors.append("Password is required.")
                if new_password != confirm_pw: errors.append("Passwords do not match.")
                if len(new_password) < 8:      errors.append("Password must be at least 8 characters.")

                if errors:
                    for e in errors: st.error(e)
                else:
                    try:
                        created = create_user(
                            firm_id    = c["firm_id"],
                            email      = new_email.strip(),
                            password   = new_password,
                            full_name  = new_full_name.strip(),
                            role       = new_role_sel,
                            created_by = c["user_id"],
                        )
                        log_action(c["firm_id"], c["user_id"], "user_create",
                                   {"new_user_email": new_email, "role": new_role_sel})
                        st.success(f"✅ User **{new_full_name}** ({new_email}) created as **{new_role_sel}**.")
                        st.info("Share the temporary password with them securely. They can change it after logging in.")
                    except ValueError as e:
                        st.error(str(e))


# ── Document helpers ──────────────────────────────────────────────────────────
def extract_text_pdf(b):
    if not PYPDF2_OK: return "","empty"
    try:
        pages = [p.extract_text() or "" for p in PyPDF2.PdfReader(io.BytesIO(b)).pages]
        t = "\n\n".join(pages).strip()
        return t, ("text" if len(t)>100 else "empty")
    except: return "","empty"

def pdf_to_images(b):
    if not PYMUPDF_OK: return []
    doc = fitz.open(stream=b, filetype="pdf")
    imgs = []; mat = fitz.Matrix(150/72,150/72)
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        imgs.append(base64.standard_b64encode(pix.tobytes("png")).decode())
    doc.close(); return imgs

def extract_docx(b):
    if not DOCX_OK: return ""
    doc = DocxDocument(io.BytesIO(b)); parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: parts.append(f"## {t}" if p.style.name.startswith("Heading") else t)
    for table in doc.tables:
        parts.append("\n[TABLE]")
        for row in table.rows: parts.append(" | ".join(c.text.strip() for c in row.cells))
        parts.append("[END TABLE]")
    return "\n".join(parts)

def process_file(uploaded):
    name=uploaded.name; ext=Path(name).suffix.lower(); b=uploaded.read()
    r={"name":name,"ext":ext,"method":"unsupported","text":"","images":[],
       "page_count":0,"size_kb":len(b)//1024,"badge":""}
    if ext==".pdf":
        text,method=extract_text_pdf(b)
        if method=="text":
            r.update({"method":"text","text":text,"badge":'<span class="badge badge-firm">📄 TEXT PDF</span>'})
            try: r["page_count"]=len(PyPDF2.PdfReader(io.BytesIO(b)).pages)
            except: pass
        elif PYMUPDF_OK:
            imgs=pdf_to_images(b)
            r.update({"method":"vision","images":imgs,"page_count":len(imgs),
                       "badge":'<span class="badge badge-case">🔍 SCANNED PDF — OCR</span>'})
    elif ext in [".docx",".doc"]:
        text=extract_docx(b)
        if text: r.update({"method":"docx","text":text,
                            "badge":'<span class="badge badge-case">📝 WORD DOCUMENT</span>'})
    return r


# ── Claude calls ──────────────────────────────────────────────────────────────
def call_claude(prompt, history=None):
    msgs = (history or []) + [{"role":"user","content":prompt}]
    resp = get_anthropic().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1500,
        system=[{"type":"text","text":SYSTEM_PROMPT,"cache_control":{"type":"ephemeral"}}],
        messages=msgs)
    return resp.content[0].text

def call_claude_vision(images, prompt):
    content = []
    for i,img in enumerate(images[:20]):
        content.append({"type":"text","text":f"Page {i+1}:"})
        content.append({"type":"image","source":{"type":"base64","media_type":"image/png","data":img}})
    if len(images)>20: content.append({"type":"text","text":f"[{len(images)} pages — first 20 shown]"})
    content.append({"type":"text","text":prompt})
    resp = get_anthropic().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1500,
        system=SYSTEM_PROMPT, messages=[{"role":"user","content":content}])
    return resp.content[0].text

def review_document(doc):
    prompt = f"Review this legal document and provide a structured analysis.\n\nFilename: {doc['name']}\n"
    if doc["method"]=="vision": return call_claude_vision(doc["images"], prompt+"Pages shown above.")
    return call_claude(prompt+f"\nContent:\n{doc['text'][:14000]}")


# ════════════════════════════════════════════════════════════════
# MAIN — session gate
# ════════════════════════════════════════════════════════════════

# Initialise session state
for k,v in [("ctx",{}),("chat_history",[]),("current_doc",None),("review_result","")]:
    if k not in st.session_state: st.session_state[k] = v

# Validate existing session on every rerun
if st.session_state["ctx"]:
    from auth import validate_session
    fresh = validate_session(st.session_state["ctx"].get("token",""))
    if fresh:
        st.session_state["ctx"] = fresh   # refresh full_name/role in case they changed
    else:
        st.session_state.clear()           # session expired — force re-login
        st.rerun()

if not st.session_state["ctx"]:
    show_login()
    st.stop()

c      = ctx()
module = show_sidebar()


# ════════════════════════════════════════════════════════════════
# MODULE: CLIENT INTAKE
# ════════════════════════════════════════════════════════════════
if module == "📋 Client Intake":
    if is_readonly():
        st.warning("Your role (readonly) cannot submit intake forms.")
        st.stop()

    st.title("New Client Intake")
    st.markdown("Complete this form to begin. The assistant will triage your matter.")

    col1, col2 = st.columns(2)
    with col1:
        name    = st.text_input("Full name")
        email   = st.text_input("Email")
        phone   = st.text_input("Phone", placeholder="+94 77 xxx xxxx")
    with col2:
        area    = st.selectbox("Area of law", ["","Contract dispute","Property / conveyancing",
                    "Employment","Family law","Commercial / corporate",
                    "Intellectual property","Criminal defence","Other"])
        urgency = st.selectbox("Urgency", ["","Urgent — within 48 hours","Within the week","No immediate deadline"])
        prior   = st.selectbox("Prior legal action?", ["","No, first time",
                    "Yes — previous attorney","Yes — already in proceedings"])

    matter = st.text_area("Describe your matter", height=120,
        placeholder="Briefly describe the legal issue, parties, and any key dates…")

    if st.button("Analyse & Triage →", type="primary"):
        if not name or not area or not matter:
            st.warning("Please fill in name, area of law, and matter description.")
        else:
            with st.spinner("Analysing…"):
                prompt = f"Triage this new client intake:\n\nClient: {name}\nArea: {area}\nUrgency: {urgency or 'Not specified'}\nPrior action: {prior or 'Not specified'}\nMatter:\n\"{matter}\"\n\nProvide a full intake triage report."
                result = call_claude(prompt)
                log_action(c["firm_id"], c["user_id"], "intake", {"client_name":name,"area":area})
            st.divider()
            st.subheader("Intake Analysis")
            st.caption("AI generated — attorney review required")
            st.markdown(result)


# ════════════════════════════════════════════════════════════════
# MODULE: LEGAL Q&A
# ════════════════════════════════════════════════════════════════
elif module == "💬 Legal Q&A":
    st.title("Legal Q&A")

    if not st.session_state.chat_history:
        with st.chat_message("assistant", avatar="⚖️"):
            st.markdown(f"""Good day, {c['full_name'].split()[0]}. I'm AltoLex, your legal information assistant.

I can help you understand legal concepts, explain procedures, and answer general questions about Sri Lankan and common law matters.

*I provide legal information only — not legal advice. Please consult a qualified attorney for specific situations.*

What would you like to know?""")

    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"], avatar="⚖️" if msg["role"]=="assistant" else "👤"):
            st.markdown(msg["content"])

    if user_input := st.chat_input("Ask a legal question…"):
        st.session_state.chat_history.append({"role":"user","content":user_input})
        with st.chat_message("user", avatar="👤"): st.markdown(user_input)
        with st.chat_message("assistant", avatar="⚖️"):
            with st.spinner("Thinking…"):
                history  = [{"role":m["role"],"content":m["content"]}
                             for m in st.session_state.chat_history[:-1][-10:]]
                response = call_claude(user_input, history)
                log_action(c["firm_id"], c["user_id"], "query", {"len":len(user_input)})
            st.markdown(response)
            st.session_state.chat_history.append({"role":"assistant","content":response})

    if st.session_state.chat_history:
        if st.button("Clear conversation"):
            st.session_state.chat_history = []; st.rerun()


# ════════════════════════════════════════════════════════════════
# MODULE: DOCUMENT REVIEW
# ════════════════════════════════════════════════════════════════
elif module == "📄 Document Review":
    st.title("Document Review")
    st.markdown("Upload a legal document for AI-assisted review.")

    uploaded = st.file_uploader("Upload document", type=["pdf","docx","doc"],
                                 label_visibility="collapsed")
    if uploaded:
        with st.spinner(f"Processing {uploaded.name}…"):
            doc = process_file(uploaded)
            st.session_state.current_doc = doc

        col_info, col_btn = st.columns([3,1])
        with col_info:
            st.markdown(doc["badge"], unsafe_allow_html=True)
            parts = [f"**{doc['name']}**", f"{doc['size_kb']} KB"]
            if doc["page_count"]: parts.append(f"{doc['page_count']} pages")
            st.markdown(" · ".join(parts))
            if doc["method"]=="vision":
                st.markdown(f'<div style="font-size:11.5px;color:#2d6a4f;margin-top:4px">💡 Scanned PDF — ~${doc["page_count"]*0.006:.3f} estimated cost ({doc["page_count"]} pages)</div>', unsafe_allow_html=True)
        with col_btn:
            analyse = st.button("✦ Analyse", type="primary", use_container_width=True,
                                disabled=is_readonly())

        if doc["method"]=="unsupported":
            st.error("File could not be processed.")
        else:
            col_preview, col_analysis = st.columns([1,1], gap="medium")
            with col_preview:
                st.markdown("**Preview**")
                if doc["method"]=="vision" and doc["images"]:
                    for i,img in enumerate(doc["images"][:5]):
                        st.image(base64.b64decode(img), use_container_width=True, caption=f"Page {i+1}")
                    if doc["page_count"]>5: st.caption(f"Showing first 5 of {doc['page_count']} pages")
                elif doc["text"]:
                    preview = doc["text"][:3000]+("\n\n[…truncated]" if len(doc["text"])>3000 else "")
                    st.text_area("", value=preview, height=520, disabled=True, label_visibility="collapsed")
            with col_analysis:
                st.markdown("**AI Analysis**")
                st.caption("Attorney sign-off required on all outputs")
                if analyse or st.session_state.review_result:
                    if analyse:
                        with st.spinner("Reviewing document…"):
                            result = review_document(doc)
                            st.session_state.review_result = result
                            log_action(c["firm_id"], c["user_id"], "review",
                                       {"filename":doc["name"],"method":doc["method"]})
                    if st.session_state.review_result:
                        st.markdown(st.session_state.review_result)
                        st.divider()
                        st.download_button("⬇ Download analysis",
                            data=st.session_state.review_result,
                            file_name=f"altolex_review_{Path(doc['name']).stem}.txt",
                            mime="text/plain")
                else:
                    st.info("Click **✦ Analyse** to receive a structured legal review.")
    else:
        st.markdown("""
        <div style="border:2px dashed rgba(184,147,90,0.3);border-radius:14px;padding:50px;
                    text-align:center;background:white;margin-top:16px">
            <div style="font-size:48px;margin-bottom:12px">📄</div>
            <div style="font-family:'Playfair Display',serif;font-size:18px;color:#1a1a2e;margin-bottom:8px">Upload a document to begin</div>
            <div style="font-size:13px;color:#6b6b80;line-height:1.7">Text PDFs · Scanned PDFs (OCR) · Word documents (.docx)</div>
        </div>""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# MODULE: ADMIN
# ════════════════════════════════════════════════════════════════
elif module == "⚙ Admin":
    if not is_admin():
        st.error("Access denied — admin role required.")
        st.stop()
    show_admin()
